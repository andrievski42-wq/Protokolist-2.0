from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.meeting import MeetingData


def export_word(target: Path, meeting: MeetingData) -> Path:
    doc = Document()

    title = doc.add_heading("ПРОТОКОЛ СОВЕЩАНИЯ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Наименование: {meeting.title or '—'}")
    doc.add_paragraph(f"Председатель: {meeting.chairman or '—'}")
    doc.add_paragraph(f"Секретарь: {meeting.secretary or '—'}")
    doc.add_paragraph(f"Место проведения: {meeting.place or '—'}")
    doc.add_paragraph(f"Дата создания: {meeting.created_at or '—'}")

    doc.add_heading("Краткое резюме", level=2)
    doc.add_paragraph(meeting.summary.strip() or "—")

    doc.add_heading("Присутствовали", level=2)
    doc.add_paragraph(meeting.participants.strip() or "—")

    doc.add_heading("Повестка", level=2)
    doc.add_paragraph(meeting.agenda.strip() or "—")

    doc.add_heading("Ход обсуждения", level=2)
    doc.add_paragraph(meeting.discussion.strip() or meeting.transcript.strip() or "—")

    doc.add_heading("Принятые решения", level=2)
    doc.add_paragraph(meeting.decisions.strip() or "—")

    doc.add_heading("Поручения", level=2)
    if meeting.tasks:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Ответственный"
        headers[1].text = "Поручение"
        headers[2].text = "Срок"
        headers[3].text = "Статус"

        for task in meeting.tasks:
            row = table.add_row().cells
            row[0].text = task.responsible
            row[1].text = task.description
            row[2].text = task.deadline
            row[3].text = task.status
    else:
        doc.add_paragraph("Поручения не внесены.")

    doc.add_heading("Полная стенограмма", level=2)
    doc.add_paragraph(meeting.transcript.strip() or "—")

    doc.save(target)
    return target
